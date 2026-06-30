"""
PDF-to-docx reconstruction for retrospective PDFs with no source document.

Rebuilds a properly structured Word document from a docling extraction,
which can then be re-exported as a fully tagged PDF.
"""

from __future__ import annotations

import logging
import os
import re
import tempfile
from pathlib import Path

from core._image_helpers import _crop_region_as_image, _extract_page_images
from core.models import AuditReport

logger = logging.getLogger(__name__)


def _infer_heading_level(text: str) -> int:
    """
    Infer heading level from common academic numbering conventions and casing.

    Rules (in priority order):
      Roman numeral prefix  → H1  (e.g. "I. Introduction", "IV Results")
      Capital letter prefix → H2  (e.g. "A. Background")
      Digit prefix          → H3  (e.g. "1. Method")
      ALL-CAPS short text   → H1  (e.g. "ABSTRACT", "RESULTS")
      Anything else         → H2  (safer default than H1 — prevents body text
                                    mislabelled as section_header from dominating
                                    document structure)
    """
    t = text.strip()
    if re.match(r"^(I|II|III|IV|V|VI|VII|VIII|IX|X)[\.\s]", t, re.IGNORECASE):
        return 1
    if re.match(r"^[A-Z]\.\s", t):
        return 2
    if re.match(r"^\d+\.\s", t):
        return 3
    if t.isupper() and len(t.split()) <= 12:
        return 1
    return 2


def _calculate_image_width(element: dict):
    """
    Return a python-docx ``Inches`` width for an image element based on its
    bbox relative to a standard letter-page width (612 pts).

    The ratio is clamped to [0.2, 0.85] so images are never tiny slivers or
    bleed-edge wide. Falls back to 4.5 inches when no bbox is available.
    """
    from docx.shared import Inches
    bbox = element.get("bbox")
    if bbox and len(bbox) == 4:
        img_width_pts = bbox[2] - bbox[0]
        ratio = img_width_pts / 612.0
        ratio = max(0.2, min(ratio, 0.85))
        return Inches(6.5 * ratio)
    return Inches(4.5)


def _set_picture_alt_text(run, alt_text: str) -> None:
    """Write alt text to the drawing element's docPr XML attribute."""
    alt_text = (alt_text or "").strip()
    if not alt_text:
        return
    # Deduplicate if string is "X\nX" pattern (duplicate from double call)
    lines = alt_text.split("\n")
    if len(lines) == 2 and lines[0].strip() == lines[1].strip():
        alt_text = lines[0].strip()
    try:
        from lxml import etree  # noqa: F401
        WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        inline = run._r.find(f'.//{{{WP_NS}}}inline')
        if inline is None:
            inline = run._r.find(f'.//{{{WP_NS}}}anchor')
        if inline is not None:
            docPr = inline.find(f'{{{WP_NS}}}docPr')
            if docPr is not None:
                docPr.set('descr', alt_text)
                docPr.set('title', alt_text[:255])
    except Exception:
        pass  # fail silently, never crash rebuild


def _get_table_grid(element: dict) -> tuple[list | None, bool]:
    """
    Read cell content from the extraction dict's ``cells`` field.

    Returns ``(grid, has_header)`` where ``grid`` is a non-empty list-of-lists
    of strings, or ``None`` when no structured data is available.
    The ``cells`` field is populated by ``extractor.extract()`` /
    ``extractor.extract_docx()`` from the underlying docling or python-docx
    table representation.
    """
    has_header = element.get("has_header_row") == "true"
    cells = element.get("cells")
    if cells and isinstance(cells, list) and len(cells) > 0:
        return cells, has_header
    return None, has_header


# Bbox convention: docling returns [left, top, right, bottom] where y is
# measured from the page BOTTOM (PDF standard, not screen convention).
# So bbox[1] = top (larger y) and bbox[3] = bottom (smaller y), giving top > bottom.
def _validate_bbox(
    bbox: list,
    page_width: float,
    page_height: float,
) -> tuple[bool, str]:
    """
    Validate that a bbox is sane.
    Returns (is_valid, reason).

    Expects docling's [left, top, right, bottom] convention where y is
    measured from the page bottom, so top > bottom.
    """
    if not bbox or len(bbox) != 4:
        return False, "bbox missing or malformed"

    x0, y0, x1, y1 = bbox  # x0=left, y0=top, x1=right, y1=bottom (docling convention)

    # Width and height must be positive
    width = x1 - x0
    height = y0 - y1  # top - bottom (both measured from page bottom, top > bottom)
    if width <= 0 or height <= 0:
        return False, f"non-positive dimensions: {width}x{height}"

    # Bbox must be within page bounds (allow small margin for floating point).
    # In docling's bottom-origin system: y0=top ≤ page_height, y1=bottom ≥ 0.
    margin = 5.0
    if (x0 < -margin or y1 < -margin
            or x1 > page_width + margin
            or y0 > page_height + margin):
        return False, (
            f"bbox extends outside page bounds: "
            f"{bbox} for page "
            f"{page_width}x{page_height}"
        )

    # Aspect ratio sanity (not absurdly thin)
    aspect = max(width, height) / min(width, height)
    if aspect > 100:
        return False, f"extreme aspect ratio: {aspect:.1f}"

    # Minimum size check (10pt minimum)
    if width < 10 or height < 10:
        return False, f"too small: {width}x{height}"

    return True, "ok"


def _match_image_by_bbox(
    target_bbox: list,
    pymupdf_images: list,
    page_height: float = 792.0,
    tolerance: float = 10.0,
) -> dict | None:
    """
    Match a docling element bbox to the closest pymupdf-extracted image
    by bbox center distance. Returns None if no acceptable match found.

    target_bbox is in docling's bottom-origin [left, top, right, bottom]
    convention. pymupdf_images carry bboxes in PyMuPDF's top-origin system
    (y increases downward). page_height is used to convert between them.
    """
    if not target_bbox:
        return None

    tx = (target_bbox[0] + target_bbox[2]) / 2
    # Convert docling bottom-origin centre-y to PyMuPDF top-origin centre-y
    ty = page_height - (target_bbox[1] + target_bbox[3]) / 2

    best_match = None
    best_distance = float('inf')

    for img in pymupdf_images:
        img_bbox = img.get("bbox")
        if not img_bbox:
            continue
        ix = (img_bbox[0] + img_bbox[2]) / 2
        iy = (img_bbox[1] + img_bbox[3]) / 2

        distance = (
            (tx - ix) ** 2 + (ty - iy) ** 2
        ) ** 0.5

        if distance < best_distance:
            best_distance = distance
            best_match = img

    # Only return match if center distance is within reasonable tolerance
    # relative to bbox size
    if best_match and best_distance < (
        max(
            target_bbox[2] - target_bbox[0],
            target_bbox[3] - target_bbox[1],
        ) / 2 + tolerance
    ):
        return best_match

    return None


def _image_phash(img_bytes: bytes) -> str | None:
    """
    Compute a simple perceptual hash for image comparison.
    Returns hex string or None on failure.
    """
    try:
        import io
        from PIL import Image
        img = Image.open(io.BytesIO(img_bytes))
        img = img.convert("L").resize((16, 16), Image.LANCZOS)
        pixels = list(img.getdata())
        avg = sum(pixels) / len(pixels)
        bits = ["1" if p > avg else "0" for p in pixels]
        bit_string = "".join(bits)
        return hex(int(bit_string, 2))[2:]
    except Exception:
        return None


def _hamming_distance(h1: str, h2: str) -> int:
    """Compute hamming distance between two hex hashes."""
    if not h1 or not h2 or len(h1) != len(h2):
        return 999
    try:
        n1 = int(h1, 16)
        n2 = int(h2, 16)
        return bin(n1 ^ n2).count("1")
    except Exception:
        return 999


def rebuild_as_docx(
    extraction: dict,
    audit_report: AuditReport,
    user_inputs: dict,
    output_path: str,
    pdf_path: str | None = None,
    approved_fixes: list[dict] | None = None,
) -> tuple[str, list[dict]]:
    """
    Reconstruct a structured Word document from a PDF extraction.

    Parameters: extraction (docling dict), audit_report (metadata fixes),
    user_inputs (element_id → alt text), output_path (.docx path),
    pdf_path (source PDF for image/formula/table extraction),
    approved_fixes (Stage 3 fixes).
    Returns (output_path, extraction_issues) after writing.
    """
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = DocxDocument()

    # ── Extraction quality tracking ───────────────────────────────────────
    extraction_issues: list[dict] = []

    # ── Build page dimension lookup for bbox validation ───────────────────
    _page_dims: dict[int, tuple[float, float]] = {}
    if pdf_path:
        try:
            import fitz  # type: ignore
            _fitz_doc = fitz.open(pdf_path)
            for _pno in range(_fitz_doc.page_count):
                _fitz_page = _fitz_doc[_pno]
                _page_dims[_pno + 1] = (
                    _fitz_page.rect.width,
                    _fitz_page.rect.height,
                )
            _fitz_doc.close()
        except Exception:
            pass

    # ── Build O(1) lookup for approved fixes keyed by element_id ─────────
    fixes_by_element_id: dict[str, dict] = {
        f["element_id"]: f
        for f in (approved_fixes or [])
        if f.get("element_id")
    }

    # ── Document metadata ─────────────────────────────────────────────────
    title = extraction.get("metadata", {}).get("title") or ""
    # Discard bare filenames set by the extractor (temp paths have no spaces).
    # A real document title will almost always contain at least one space.
    if title and " " not in title:
        title = ""
    language = "en-US"
    for fix in (audit_report.metadata_fixes or []):
        if fix.get("field") == "language" and fix.get("value"):
            language = fix["value"]
        if fix.get("field") == "title" and fix.get("value"):
            title = fix["value"]

    doc.core_properties.title = title
    doc.core_properties.language = language

    # ── Pre-scan: find title element to hoist to top of document ─────────
    # Pass 1: explicit "title" docling label (any page, first occurrence)
    _title_element: dict | None = None
    for _pg in extraction.get("pages", []):
        for _el in _pg.get("elements", []):
            if _el.get("docling_label") == "title" and not _title_element:
                _title_element = _el
                break
        if _title_element:
            break

    # Pass 2: all-caps section_header on page 1 without a numbered prefix.
    # Docling frequently labels document titles as section_header when the
    # PDF has no explicit title tag. Criteria: all-uppercase, ≥ 4 words,
    # no Roman-numeral / capital-letter / digit section number at the start.
    if not _title_element and extraction.get("pages"):
        for _el in extraction["pages"][0].get("elements", []):
            if _el.get("docling_label") == "section_header":
                _sh_text = (_el.get("text") or "").strip()
                if (
                    _sh_text.isupper()
                    and len(_sh_text.split()) >= 4
                    and not re.match(r"^(I|II|III|IV|V|VI|VII|VIII|IX|X)[\.\s]", _sh_text, re.IGNORECASE)
                    and not re.match(r"^[A-Z]\.\s", _sh_text)
                    and not re.match(r"^\d+\.\s", _sh_text)
                ):
                    _title_element = _el
                    break

    # Pass 3: page_header on page 1 that is not a bare page number
    if not _title_element and extraction.get("pages"):
        for _el in extraction["pages"][0].get("elements", []):
            if _el.get("docling_label") == "page_header":
                _hdr_text = (_el.get("text") or "").strip()
                if (
                    not re.fullmatch(r"Page\s*\d+|\d+", _hdr_text, re.IGNORECASE)
                    and len(_hdr_text) > 5
                ):
                    _title_element = _el
                    break

    # ── State for page-header deduplication and title detection ──────────
    title_added: bool = False
    seen_page_headers: set[str] = set()

    # Write hoisted title as the very first paragraph
    if _title_element:
        _hoisted_text = (_title_element.get("text") or "").strip()
        if _hoisted_text:
            doc.add_heading(_hoisted_text, level=0)
            title_added = True

    # ── Flatten all elements with page number for look-ahead ─────────────
    all_elements: list[tuple[int, dict]] = []
    for page in extraction.get("pages", []):
        pno = page.get("page_number", 1)
        for el in page.get("elements", []):
            all_elements.append((pno, el))

    for i, (page_no, element) in enumerate(all_elements):
        # Skip the element that was already hoisted as title
        if (
            _title_element is not None
            and title_added
            and element.get("id") == _title_element.get("id")
        ):
            continue
        next_el = all_elements[i + 1][1] if i + 1 < len(all_elements) else None

        label = element.get("docling_label", "text")
        text = (element.get("text") or "").strip()
        element_id = element.get("id", "")
        elem_bbox = element.get("bbox")

        # ── Page footer — always skip ─────────────────────────────────
        if label == "page_footer":
            continue

        # ── Page header — preserve page-1 identifier; skip running headers ─
        elif label == "page_header":
            if not text:
                continue
            # Skip bare page numbers: "3", "Page 3", "Page  3", etc.
            if re.fullmatch(r"Page\s*\d+|\d+", text, re.IGNORECASE):
                continue
            if text in seen_page_headers:
                continue
            seen_page_headers.add(text)
            # Page 1 only: include as italic centered identifier (e.g. "Experiment IV-A")
            if page_no == 1:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.italic = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Pages 2+: skip — it's a running header, not body content

        # ── Title ─────────────────────────────────────────────────────
        elif label == "title":
            if not text:
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.name = "Times New Roman"
            title_added = True

        # ── Section header ────────────────────────────────────────────
        elif label == "section_header":
            if not text:
                continue
            fix = fixes_by_element_id.get(element_id, {})
            level = (
                int(fix["value"])
                if fix.get("value") and str(fix["value"]).isdigit()
                else _infer_heading_level(text)
            )
            h = doc.add_heading(text, level=level)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)

        # ── Body text / paragraph ─────────────────────────────────────
        elif label in ("text", "paragraph"):
            if not text:
                continue
            next_label = next_el.get("docling_label", "") if next_el else ""
            is_likely_heading = (
                re.match(r"^(I|II|III|IV|V|VI|VII|VIII|IX|X)[\.\s]", text, re.IGNORECASE)
                or re.match(r"^[A-Z]\.\s", text)
                or (text.isupper() and len(text.split()) <= 12)
                or (len(text.split()) <= 7 and next_label in ("text", "paragraph", "list_item"))
            )
            if is_likely_heading:
                level = _infer_heading_level(text)
                h = doc.add_heading(text, level=level)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Inches(0.5)
                run = p.add_run(text)
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)

        # ── Caption ───────────────────────────────────────────────────
        elif label == "caption":
            if not text:
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.italic = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

        # ── Footnote ──────────────────────────────────────────────────
        elif label == "footnote":
            if not text:
                continue
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)

        # ── List item ─────────────────────────────────────────────────
        elif label == "list_item":
            if not text:
                continue
            doc.add_paragraph(text, style="List Bullet")

        # ── Code ──────────────────────────────────────────────────────
        elif label == "code":
            if not text:
                continue
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = "Courier New"
            run.font.size = Pt(10)

        # ── Formula ───────────────────────────────────────────────────
        elif label == "formula":
            _fix = fixes_by_element_id.get(element_id, {})
            alt_text = _fix.get("user_value") or _fix.get("value") or user_inputs.get(element_id, "")
            rendered = False
            if pdf_path and elem_bbox:
                _pw, _ph = _page_dims.get(page_no, (612.0, 792.0))
                _bbox_valid, _bbox_reason = _validate_bbox(elem_bbox, _pw, _ph)
                if not _bbox_valid:
                    _ph_para = doc.add_paragraph()
                    _ph_run = _ph_para.add_run(
                        f"[Formula on page {page_no} could not be extracted "
                        f"reliably: {_bbox_reason}. Please refer to the original PDF.]"
                    )
                    _ph_run.italic = True
                    _ph_run.font.color.rgb = RGBColor(0xCC, 0x44, 0x00)
                    extraction_issues.append({
                        "type": "image_bbox_invalid",
                        "page": page_no,
                        "element_id": element.get("id"),
                        "reason": _bbox_reason,
                    })
                    continue
                img_bytes = _crop_region_as_image(pdf_path, page_no - 1, elem_bbox, dpi=200)
                if img_bytes:
                    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
                    tmp.write(img_bytes)
                    tmp.close()
                    try:
                        p_formula = doc.add_paragraph()
                        p_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_formula = p_formula.add_run()
                        run_formula.add_picture(tmp.name, width=_calculate_image_width(element))
                        _set_picture_alt_text(run_formula, alt_text)
                        rendered = True
                    finally:
                        os.unlink(tmp.name)
            if not rendered:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                fb = p.add_run(f"[Equation: {alt_text}]" if alt_text else "[Equation — provide plain text description]")
                fb.italic = True
                # Only add placeholder caption when image extraction failed
                cap = doc.add_paragraph()
                cap_run = cap.add_run("[Formula — edit in source document]")
                cap_run.italic = True
                cap_run.font.size = Pt(9)
                cap_run.font.color.rgb = RGBColor(128, 128, 128)

        # ── Picture ───────────────────────────────────────────────────
        elif label == "picture":
            _fix = fixes_by_element_id.get(element_id, {})
            alt_text = _fix.get("user_value") or _fix.get("value") or user_inputs.get(element_id, "")
            img_bytes: bytes | None = None

            if pdf_path and elem_bbox:
                # Validate bbox before attempting any extraction
                _pw, _ph = _page_dims.get(page_no, (612.0, 792.0))
                _bbox_valid, _bbox_reason = _validate_bbox(elem_bbox, _pw, _ph)
                if not _bbox_valid:
                    _ph_para = doc.add_paragraph()
                    _ph_run = _ph_para.add_run(
                        f"[Image on page {page_no} could not be extracted "
                        f"reliably: {_bbox_reason}. Please refer to the original PDF.]"
                    )
                    _ph_run.italic = True
                    _ph_run.font.color.rgb = RGBColor(0xCC, 0x44, 0x00)
                    extraction_issues.append({
                        "type": "image_bbox_invalid",
                        "page": page_no,
                        "element_id": element.get("id"),
                        "reason": _bbox_reason,
                    })
                    continue

                # Try spatial matching against embedded page images first
                page_images = _extract_page_images(pdf_path, page_no - 1)
                if page_images:
                    _pymupdf_imgs = [
                        {"bbox": list(k), "bytes": v}
                        for k, v in page_images.items()
                    ]
                    _ph_for_match = _page_dims.get(page_no, (612.0, 792.0))[1]
                    _best_match = _match_image_by_bbox(
                        elem_bbox, _pymupdf_imgs, page_height=_ph_for_match
                    )
                    if _best_match:
                        img_bytes = _best_match["bytes"]

                # Fall back to bbox crop if no spatial match found
                if not img_bytes:
                    img_bytes = _crop_region_as_image(pdf_path, page_no - 1, elem_bbox)

            if img_bytes:
                tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
                tmp.write(img_bytes)
                tmp.close()
                try:
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_img = p_img.add_run()
                    run_img.add_picture(tmp.name, width=_calculate_image_width(element))
                    _set_picture_alt_text(run_img, alt_text)
                    # Hash verification: compare inserted image to source region
                    if pdf_path and elem_bbox:
                        _source_crop = _crop_region_as_image(
                            pdf_path, page_no - 1, elem_bbox
                        )
                        _src_hash = _image_phash(_source_crop)
                        _ins_hash = _image_phash(img_bytes)
                        if _src_hash and _ins_hash:
                            _dist = _hamming_distance(_src_hash, _ins_hash)
                            if _dist > 50:
                                extraction_issues.append({
                                    "type": "image_hash_mismatch",
                                    "page": page_no,
                                    "element_id": element.get("id"),
                                    "reason": (
                                        f"inserted image differs from source region "
                                        f"(hamming distance: {_dist})"
                                    ),
                                })
                finally:
                    os.unlink(tmp.name)
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                placeholder = f"[Image: {alt_text}]" if alt_text else "[Image — alt text required: describe this image]"
                p.add_run(placeholder)

        # ── Table ─────────────────────────────────────────────────────
        elif label == "table":
            grid, has_header = _get_table_grid(element)

            if grid:
                n_rows = len(grid)
                n_cols = max((len(row) for row in grid), default=1)
                table = doc.add_table(rows=n_rows, cols=n_cols)
                try:
                    table.style = "Table Grid"
                except KeyError:
                    pass
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
                for row_idx, row_data in enumerate(grid):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < n_cols:
                            cell = table.cell(row_idx, col_idx)
                            cell.paragraphs[0].text = str(cell_text or "")
                            if has_header and row_idx == 0:
                                for r in cell.paragraphs[0].runs:
                                    r.bold = True
            else:
                rendered = False
                img_bytes = None
                if pdf_path and elem_bbox:
                    _tpw, _tph = _page_dims.get(page_no, (612.0, 792.0))
                    _tvalid, _treason = _validate_bbox(elem_bbox, _tpw, _tph)
                    if not _tvalid:
                        extraction_issues.append({
                            "type": "image_bbox_invalid",
                            "page": page_no,
                            "element_id": element.get("id"),
                            "reason": _treason,
                        })
                    else:
                        logger.warning("Table %s: no structured data, falling back to image crop", element_id)
                        img_bytes = _crop_region_as_image(pdf_path, page_no - 1, elem_bbox, dpi=150)
                if img_bytes:
                    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
                    tmp.write(img_bytes)
                    tmp.close()
                    try:
                        doc.add_picture(tmp.name, width=_calculate_image_width(element))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        rendered = True
                    finally:
                        os.unlink(tmp.name)
                if not rendered:
                    doc.add_paragraph("[Table — could not extract content]")
            doc.add_paragraph()  # spacing after table

        # ── Unrecognised label — treat as body paragraph ──────────────
        else:
            if not text:
                continue
            logger.warning("Unrecognised docling_label %r for element %s — treating as body paragraph", label, element_id)
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

    doc.save(output_path)
    return output_path, extraction_issues


def verify_content_fidelity(
    original_pdf_path: str,
    rebuilt_docx_path: str,
    extraction: dict,
) -> dict:
    """
    Compares body text from the original PDF extraction and the rebuilt
    Word document. Returns a fidelity report.

    Returns:
        {
            "match_percentage": float,
            "missing_text": list[str],
            "summary": str,
        }
    """
    from docx import Document

    # Collect text blocks from extraction (already parsed from original PDF)
    _SKIP_LABELS = {"page_header", "page_footer", "footnote"}
    original_text_blocks: list[str] = []
    for page in extraction.get("pages", []):
        for el in page.get("elements", []):
            label = el.get("docling_label", "")
            if label in _SKIP_LABELS:
                continue
            text = (el.get("text") or "").strip()
            if text and len(text) > 10:
                original_text_blocks.append(text)

    # Extract text from rebuilt docx
    try:
        doc = Document(rebuilt_docx_path)
        rebuilt_text = " ".join(
            para.text.strip() for para in doc.paragraphs if para.text.strip()
        )
    except Exception:
        rebuilt_text = ""

    rebuilt_normalized = " ".join(rebuilt_text.split())

    # Check what percentage of original blocks appear in rebuilt text
    missing: list[str] = []
    matched = 0
    for block in original_text_blocks:
        test_chunk = " ".join(block.split())[:50]
        if test_chunk and test_chunk in rebuilt_normalized:
            matched += 1
        else:
            missing.append(block[:100] + ("…" if len(block) > 100 else ""))

    total = len(original_text_blocks)
    pct = (matched / total * 100) if total > 0 else 100.0

    if pct >= 95:
        summary = (
            f"Content fidelity is high. {matched} of {total} text blocks "
            "from the original document were found in the rebuilt output."
        )
    elif pct >= 80:
        summary = (
            f"Content fidelity is acceptable but {total - matched} text blocks "
            "could not be matched. Review the rebuilt document carefully."
        )
    else:
        summary = (
            f"Content fidelity is low. {total - matched} of {total} text blocks "
            "could not be matched. The rebuilt document may have significant content gaps."
        )

    return {
        "match_percentage": round(pct, 1),
        "missing_text": missing[:10],
        "summary": summary,
    }
