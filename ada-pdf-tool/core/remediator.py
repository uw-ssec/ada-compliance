"""
Remediation layer — applies approved fixes to PDF (pikepdf) and docx (python-docx).

PDF: only metadata fixes (language, title) and bookmark generation can be
applied programmatically. Structural fixes require editing the source document.

docx: all structural fix types are writable via python-docx.
"""

from __future__ import annotations

import shutil
from pathlib import Path
from typing import Any

import pikepdf


_SKIP_REASON = (
    "Requires accessibility tag tree — fix in source document (Word/Google Docs) "
    "and re-export with accessibility settings enabled."
)

_STRUCTURAL_FIX_TYPES = {"heading_tag", "alt_text", "link_tag", "table_header"}


def remediate(
    input_path: str,
    output_path: str,
    approved_fixes: list,
) -> dict:
    """
    Apply approved fixes to input_path and write the result to output_path.

    Parameters
    ----------
    input_path:
        Path to the original PDF. Never modified.
    output_path:
        Path where the remediated PDF will be written.
    approved_fixes:
        List of fix dicts. Each fix must have at minimum:
          - "fix_type": one of "set_language", "set_title", "set_bookmarks",
            or a structural type (will be skipped with explanation)
          - For "set_language": no additional fields required (sets "en-US")
          - For "set_title": "value" str with the title
          - For "set_bookmarks": "headings" list of {"text": str, "page": int}

    Returns
    -------
    dict with keys:
        "applied": list of str — descriptions of applied fixes
        "skipped": list of str — fixes skipped with reason
        "errors":  list of str — any exceptions encountered
    """
    applied: list[str] = []
    skipped: list[str] = []
    errors: list[str] = []

    input_path = str(input_path)
    output_path = str(output_path)

    try:
        with pikepdf.open(input_path) as pdf:
            for fix in approved_fixes:
                fix_type = fix.get("fix_type", "")

                if fix_type in _STRUCTURAL_FIX_TYPES:
                    skipped.append(f"{fix_type} on element {fix.get('element_id', '?')}: {_SKIP_REASON}")
                    continue

                if fix_type == "set_language":
                    pdf.Root.Lang = pikepdf.String("en-US")
                    applied.append("Set document language to en-US")

                elif fix_type == "set_title":
                    value = fix.get("value", "")
                    if value:
                        pdf.docinfo["/Title"] = value
                        try:
                            with pdf.open_metadata() as meta:
                                meta["dc:title"] = value
                        except Exception:
                            pass  # XMP write is best-effort
                        applied.append(f"Set document title to: {value}")
                    else:
                        skipped.append("set_title: no value provided")

                elif fix_type == "set_bookmarks":
                    headings = fix.get("headings", [])
                    if headings:
                        with pdf.open_outline() as outline:
                            outline.root.clear()
                            for h in headings:
                                page_no = h.get("page", 1)
                                text = h.get("text", "")
                                if text:
                                    # pikepdf OutlineItem page numbers are 0-based
                                    item = pikepdf.OutlineItem(text, page_no - 1)
                                    outline.root.append(item)
                        pdf.Root["/PageMode"] = pikepdf.Name("/UseOutlines")
                        applied.append(f"Added {len(headings)} bookmarks from heading structure")
                    else:
                        skipped.append("set_bookmarks: no headings provided")

                else:
                    skipped.append(f"Unknown fix type '{fix_type}': {_SKIP_REASON}")

            pdf.save(output_path)

    except Exception as exc:
        errors.append(str(exc))
        # Best-effort: copy original so output_path always exists
        try:
            shutil.copy2(input_path, output_path)
        except Exception:
            pass

    return {"applied": applied, "skipped": skipped, "errors": errors}


def validate_fixes(output_path: str, applied_fixes: list) -> dict:
    """
    Open the output file and check whether specific fixes are confirmed present.

    Parameters
    ----------
    output_path:
        Path to the remediated file (PDF or docx).
    applied_fixes:
        List of fix dicts that were passed to remediate() / remediate_docx().

    Returns
    -------
    dict with keys:
        "passed":  list[str] — confirmed present
        "failed":  list[str] — expected but not found
        "skipped": list[str] — could not check
    """
    passed: list[str] = []
    failed: list[str] = []
    skipped: list[str] = []

    suffix = Path(output_path).suffix.lower()

    fix_types = {f.get("fix_type") for f in applied_fixes}

    # ── PDF checks (pikepdf) ──────────────────────────────────────────────
    if suffix == ".pdf":
        try:
            with pikepdf.open(output_path) as pdf:
                if "set_language" in fix_types:
                    lang = pdf.Root.get("/Lang")
                    if lang is not None and str(lang).strip():
                        passed.append(f"Language metadata: {str(lang)} confirmed")
                    else:
                        failed.append("Language metadata: not found in output")

                if "set_title" in fix_types:
                    title_val = pdf.docinfo.get("/Title")
                    if title_val is not None and str(title_val).strip():
                        passed.append(f"Document title: '{str(title_val)}' confirmed")
                    else:
                        failed.append("Document title: not found in output")

                if "set_bookmarks" in fix_types:
                    try:
                        outlines = pdf.Root.get("/Outlines")
                        if outlines is not None:
                            kids = outlines.get("/Count")
                            count = int(str(kids)) if kids is not None else 0
                            if count > 0:
                                passed.append(f"Bookmarks: {count} entries confirmed")
                            else:
                                # Count may be absent; check for /First child
                                first = outlines.get("/First")
                                if first is not None:
                                    passed.append("Bookmarks: entries confirmed")
                                else:
                                    failed.append("Bookmarks: outline missing from output")
                        else:
                            failed.append("Bookmarks: outline missing from output")
                    except Exception:
                        skipped.append("Bookmarks: could not verify outline structure")
        except Exception as exc:
            skipped.append(f"PDF validation skipped: {exc}")

    # ── docx checks (python-docx) ─────────────────────────────────────────
    elif suffix == ".docx":
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(output_path)
            _doc_modified = False

            if "set_alt_text" in fix_types:
                alt_fixes = [f for f in applied_fixes if f.get("fix_type") == "set_alt_text"]
                for fix in alt_fixes:
                    eid = fix.get("element_id", "?")
                    para_idx = fix.get("paragraph_index")
                    found = False
                    paragraphs_to_check = []
                    if para_idx is not None:
                        try:
                            pi = int(para_idx)
                            if 0 <= pi < len(doc.paragraphs):
                                paragraphs_to_check = [doc.paragraphs[pi]]
                        except (TypeError, ValueError):
                            pass
                    if not paragraphs_to_check:
                        paragraphs_to_check = list(doc.paragraphs)
                    for p in paragraphs_to_check:
                        for run in p.runs:
                            for drawing in run._r.iter():
                                tag = getattr(drawing, "tag", "") or ""
                                if "docPr" in tag:
                                    descr = drawing.get("descr", "")
                                    if descr and descr.strip():
                                        found = True
                    if found:
                        passed.append(f"Alt text: confirmed on {eid}")
                    else:
                        failed.append(f"Alt text: not found on {eid}")

            # For rebuilt Word docs — check heading styles present
            if "set_heading_style" in fix_types or not fix_types.intersection(
                {"set_language", "set_title", "set_bookmarks", "set_alt_text", "set_table_header"}
            ):
                has_heading = any(
                    (p.style.name or "").startswith("Heading")
                    for p in doc.paragraphs
                    if p.style
                )
                if has_heading:
                    passed.append("Heading styles: present in rebuilt document")
                else:
                    failed.append("Heading styles: none found in rebuilt document")

            # ── Alt text deduplication scan (safety net, always runs) ─────────
            # Detects and fixes duplication patterns regardless of what fixes were
            # applied, so downstream failures never produce doubled alt text.
            _WP_NS = (
                "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
            )
            for _para in doc.paragraphs:
                for _run in _para.runs:
                    _drawings = _run._r.findall(f".//{{{_WP_NS}}}docPr")
                    for _docPr in _drawings:
                        _descr = _docPr.get("descr", "")
                        if not _descr:
                            continue

                        # Pattern 1: "X\nX" — two identical lines
                        _lines = [ln.strip() for ln in _descr.split("\n") if ln.strip()]
                        if len(_lines) == 2 and _lines[0] == _lines[1]:
                            _docPr.set("descr", _lines[0])
                            _doc_modified = True
                            passed.append(f"Alt text deduplicated: {_lines[0][:60]}")
                            continue

                        # Pattern 2: "XX" — string is exactly two copies concatenated
                        _half = len(_descr) // 2
                        if len(_descr) > 20 and _descr[:_half] == _descr[_half:]:
                            _docPr.set("descr", _descr[:_half])
                            _doc_modified = True
                            passed.append(f"Alt text deduplicated: {_descr[:_half][:60]}")
                            continue

                        # No duplication detected
                        passed.append(f"Alt text verified: {_descr[:60]}")

            if _doc_modified:
                doc.save(output_path)

        except Exception as exc:
            skipped.append(f"docx validation skipped: {exc}")

    else:
        skipped.append(f"Validation not supported for file type: {suffix}")

    return {"passed": passed, "failed": failed, "skipped": skipped}


def remediate_docx(
    input_path: str,
    output_path: str,
    approved_fixes: list,
    user_inputs: dict,
) -> dict:
    """
    Apply approved fixes to a Word .docx file and write the result to output_path.

    All fix types are writable in docx: set_language, set_title,
    set_heading_style, set_alt_text, set_table_header.

    Parameters
    ----------
    input_path:
        Path to the original .docx. Never modified.
    output_path:
        Path where the remediated .docx will be written.
    approved_fixes:
        List of fix dicts with at minimum a "fix_type" key.
    user_inputs:
        Mapping of element_id → user-provided value (e.g. alt text).

    Returns
    -------
    dict with keys "applied", "skipped", "errors".
    """
    from docx import Document as DocxDocument

    applied: list[str] = []
    skipped: list[str] = []
    errors: list[str] = []

    input_path = str(input_path)
    output_path = str(output_path)

    shutil.copy2(input_path, output_path)

    try:
        doc = DocxDocument(output_path)

        for fix in approved_fixes:
            fix_type = fix.get("fix_type", "")
            element_id = fix.get("element_id", "")

            try:
                if fix_type == "set_language":
                    doc.core_properties.language = "en-US"
                    applied.append("Set document language to en-US")

                elif fix_type == "set_title":
                    value = fix.get("value", "")
                    if value:
                        doc.core_properties.title = value
                        applied.append(f"Set document title to: {value}")
                    else:
                        skipped.append("set_title: no value provided")

                elif fix_type == "set_heading_style":
                    # Locate paragraph by paragraph_index (direct index into
                    # doc.paragraphs), falling back to text-content match.
                    para = None
                    para_idx = fix.get("paragraph_index")
                    if para_idx is not None:
                        try:
                            para_idx = int(para_idx)
                            if 0 <= para_idx < len(doc.paragraphs):
                                para = doc.paragraphs[para_idx]
                        except (TypeError, ValueError):
                            pass
                    if para is None:
                        target_text = (fix.get("text") or "").strip()
                        for p in doc.paragraphs:
                            if p.text.strip() == target_text:
                                para = p
                                break
                    if para is None:
                        skipped.append(
                            f"set_heading_style: could not locate element {element_id}"
                        )
                    else:
                        # Level from fix["value"] if it's a digit string, else default 1
                        level_raw = str(fix.get("value") or "1")
                        level = int(level_raw) if level_raw.isdigit() and 1 <= int(level_raw) <= 6 else 1
                        style_name = f"Heading {level}"
                        try:
                            para.style = doc.styles[style_name]
                        except KeyError:
                            para.style = doc.styles["Heading 1"]
                        applied.append(f"Set {style_name}: {para.text[:50]}")

                elif fix_type == "set_alt_text":
                    # Pre-resolved value in fix dict takes priority over user_inputs.
                    alt_text = fix.get("user_value") or fix.get("value") or user_inputs.get(element_id, "")
                    if not alt_text:
                        skipped.append(f"set_alt_text on {element_id}: no alt text value provided")
                        continue
                    # Use paragraph_index to target the specific paragraph when available.
                    para_idx = fix.get("paragraph_index")
                    paragraphs_to_check = []
                    if para_idx is not None:
                        try:
                            para_idx = int(para_idx)
                            if 0 <= para_idx < len(doc.paragraphs):
                                paragraphs_to_check = [doc.paragraphs[para_idx]]
                        except (TypeError, ValueError):
                            pass
                    if not paragraphs_to_check:
                        paragraphs_to_check = list(doc.paragraphs)
                    set_count = 0
                    for p in paragraphs_to_check:
                        for run in p.runs:
                            for drawing in run._r.iter():
                                tag = getattr(drawing, "tag", "") or ""
                                if "docPr" in tag:
                                    drawing.set("descr", alt_text)
                                    set_count += 1
                    if set_count:
                        applied.append(f"Set alt text on image {element_id}: {alt_text[:50]}")
                    else:
                        skipped.append(f"set_alt_text on {element_id}: no image drawing element found")

                elif fix_type == "set_table_header":
                    # Use table_index from the fix dict (set by extractor and carried
                    # through app.py). Default to 0 only as a last resort.
                    table_index = 0
                    raw_tidx = fix.get("table_index")
                    if raw_tidx is not None:
                        try:
                            table_index = int(raw_tidx)
                        except (TypeError, ValueError):
                            pass
                    if not doc.tables:
                        skipped.append(f"set_table_header on {element_id}: no tables found in document")
                    elif table_index >= len(doc.tables):
                        skipped.append(
                            f"set_table_header on {element_id}: table_index {table_index} out of range"
                        )
                    else:
                        table = doc.tables[table_index]
                        header_style = None
                        try:
                            header_style = doc.styles["Table Header"]
                        except KeyError:
                            header_style = doc.styles["Normal"]
                        for cell in table.rows[0].cells:
                            for p in cell.paragraphs:
                                p.style = header_style
                                if header_style.name == "Normal":
                                    for run in p.runs:
                                        run.bold = True
                        applied.append(f"Applied header style to first row of table {element_id}")

                else:
                    skipped.append(f"fix type '{fix_type}' not applicable to Word documents")

            except Exception as exc:
                errors.append(f"{fix_type} on {element_id}: {exc}")

        doc.save(output_path)

    except Exception as exc:
        errors.append(str(exc))
        try:
            shutil.copy2(input_path, output_path)
        except Exception:
            pass

    return {"applied": applied, "skipped": skipped, "errors": errors}
