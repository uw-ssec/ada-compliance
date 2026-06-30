"""
PDF extraction layer using docling.

Extracts structured content (text, images, tables) from programmatic PDFs
and returns a JSON-serialisable dict matching the ExtractionOutput schema.
Scanned PDFs (no embedded text) raise ValueError.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docling.document_converter import DocumentConverter, PdfFormatOption
from docling.datamodel.pipeline_options import PdfPipelineOptions
from docling.datamodel.base_models import InputFormat
from docling.datamodel.document import DoclingDocument
from docling_core.types.doc import DocItemLabel, ImageRefMode


def _make_id(counter: int) -> str:
    return f"el_{counter:03d}"


def _docling_bbox_to_pymupdf(bbox: list, page_height: float) -> tuple:
    """
    Convert a docling bbox to PyMuPDF coordinate space.

    Docling returns bboxes as [l, t, r, b] from ``_bbox_to_list`` where
    y is measured from the **bottom** of the page (PDF standard).
    ``t`` is the **top** edge (larger y, further from page bottom) and
    ``b`` is the **bottom** edge (smaller y, closer to page bottom).
    PyMuPDF measures y from the **top** of the page (y increases downward),
    so: pymupdf_y0 = page_height - t  and  pymupdf_y1 = page_height - b,
    giving y0 < y1 as required for a valid non-empty Rect.

    Parameters
    ----------
    bbox:
        [x0, t, x1, b] in docling/PDF bottom-left coords, where t > b.
    page_height:
        Height of the page in PDF points (``page.rect.height``).

    Returns
    -------
    tuple — (x0, y0, x1, y1) ready for ``fitz.Rect``.
    """
    x0, t, x1, b = bbox
    pymupdf_y0 = page_height - t  # top edge from page top (smaller value)
    pymupdf_y1 = page_height - b  # bottom edge from page top (larger value)
    return (x0, pymupdf_y0, x1, pymupdf_y1)


def _bbox_to_list(bbox: Any) -> list | None:
    """Convert a docling BoundingBox to [x0, y0, x1, y1] or return None."""
    if bbox is None:
        return None
    try:
        return [bbox.l, bbox.t, bbox.r, bbox.b]
    except AttributeError:
        return None


def _build_page_figure_alt_map(pdf_path: str) -> dict[int, list[bool]]:
    """
    Walk the PDF struct tree and return, per page, an ordered list of booleans
    indicating whether each /Figure element has a non-empty /Alt attribute.

    Returns {} when no struct tree is present or on any error.
    Page numbers are 1-indexed (matching docling provenance).
    """
    result: dict[int, list[bool]] = {}
    try:
        import pikepdf as _pike
        with _pike.open(str(pdf_path)) as pdf:
            struct_root = pdf.Root.get("/StructTreeRoot")
            if struct_root is None:
                return result

            # Build {id(page.obj): 1-indexed page number}
            page_num_map: dict[int, int] = {}
            for i, page in enumerate(pdf.pages):
                page_num_map[id(page.obj)] = i + 1

            def _walk(node: Any) -> None:  # noqa: ANN001
                if node is None:
                    return
                try:
                    node_type = node.get("/S")
                except Exception:
                    return

                if node_type is not None and str(node_type) == "/Figure":
                    try:
                        pg_ref = node.get("/Pg")
                        if pg_ref is not None:
                            pno = page_num_map.get(id(pg_ref.obj))
                            if pno is not None:
                                alt = node.get("/Alt")
                                has_alt = alt is not None and str(alt).strip() != ""
                                result.setdefault(pno, []).append(has_alt)
                                return  # do not recurse into figure children
                    except Exception:
                        pass

                # Recurse into /K children
                try:
                    kids = node.get("/K")
                    if kids is None:
                        return
                    if hasattr(kids, "__iter__") and not isinstance(kids, (str, bytes)):
                        for kid in kids:
                            try:
                                if hasattr(kid, "get"):
                                    _walk(kid)
                            except Exception:
                                pass
                    elif hasattr(kids, "get"):
                        _walk(kids)
                except Exception:
                    pass

            _walk(struct_root)
    except Exception:
        pass
    return result


def extract(pdf_path: str | Path) -> dict:
    """
    Extract structured content from a programmatic PDF.

    Parameters
    ----------
    pdf_path:
        Path to the PDF file.

    Returns
    -------
    dict
        JSON-serialisable dict matching the ExtractionOutput schema.

    Raises
    ------
    FileNotFoundError
        If the PDF does not exist at the given path.
    ValueError
        If the PDF appears to be scanned (no embedded text found).
    """
    pdf_path = Path(pdf_path)
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF not found: {pdf_path}")

    pipeline_options = PdfPipelineOptions()
    pipeline_options.do_ocr = False          # disable OCR — programmatic PDFs only
    pipeline_options.do_table_structure = True
    pipeline_options.generate_picture_images = False

    converter = DocumentConverter(
        format_options={
            InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
        }
    )

    result = converter.convert(str(pdf_path))
    doc: DoclingDocument = result.document

    # ── scanned-PDF guard ──────────────────────────────────────────────────
    # Collect all text across all pages; if completely empty, treat as scanned.
    all_text = doc.export_to_text().strip()
    if not all_text:
        raise ValueError(
            "This PDF appears to be scanned. Only programmatic PDFs with "
            "embedded text are supported."
        )

    # ── build per-page figure alt-text map (requires pikepdf, best-effort) ──
    _alt_map = _build_page_figure_alt_map(str(pdf_path))

    # ── metadata ──────────────────────────────────────────────────────────
    page_count = len(doc.pages) if doc.pages else 0

    title: str | None = None
    language: str | None = None

    # DoclingDocument stores metadata in doc.name and doc.origin;
    # use getattr throughout to be safe across docling versions.
    title = getattr(doc, "name", None) or None
    lang_raw = getattr(doc, "lang", None)
    if lang_raw:
        # lang may be a list or a string depending on docling version
        if isinstance(lang_raw, list):
            language = lang_raw[0] if lang_raw else None
        else:
            language = str(lang_raw)

    metadata = {
        "title": title,
        "language": language,
        "page_count": page_count,
    }

    # ── build a page_number → page_index map ─────────────────────────────
    # doc.pages is a dict keyed by PageNo (1-based int-like objects)
    page_numbers = sorted(doc.pages.keys(), key=lambda p: int(p))

    pages_map: dict[int, list[dict]] = {int(p): [] for p in page_numbers}

    element_counter = 1          # sequential element ID counter
    _img_counters: dict[int, int] = {}  # page_no → images seen so far on that page
    _table_counter: int = 0       # tables seen across the whole document

    # ── iterate document items ─────────────────────────────────────────────
    for item, _level in doc.iterate_items():
        label = getattr(item, "label", None)
        if label is None:
            continue

        # Determine page number for this item
        prov = getattr(item, "prov", None)
        page_no: int | None = None
        bbox_list: list | None = None

        if prov:
            first_prov = prov[0] if isinstance(prov, (list, tuple)) and prov else prov
            raw_page = getattr(first_prov, "page_no", None)
            if raw_page is not None:
                page_no = int(raw_page)
            raw_bbox = getattr(first_prov, "bbox", None)
            bbox_list = _bbox_to_list(raw_bbox)

        if page_no is None or page_no not in pages_map:
            # Fall back to page 1 if provenance is missing
            page_no = int(page_numbers[0]) if page_numbers else 1
            if page_no not in pages_map:
                pages_map[page_no] = []

        el_id = _make_id(element_counter)
        element_counter += 1

        # ── TEXT elements ──────────────────────────────────────────────────
        if label in (
            DocItemLabel.TEXT,
            DocItemLabel.PARAGRAPH,
            DocItemLabel.SECTION_HEADER,
            DocItemLabel.TITLE,
            DocItemLabel.CAPTION,
            DocItemLabel.FOOTNOTE,
            DocItemLabel.PAGE_HEADER,
            DocItemLabel.PAGE_FOOTER,
            DocItemLabel.LIST_ITEM,
            DocItemLabel.CODE,
            DocItemLabel.FORMULA,
        ):
            text_content = getattr(item, "text", None)

            element: dict = {
                "id": el_id,
                "type": "text",
                "docling_label": label.value,
                "text": text_content,
                "font_size": None,
                "font_bold": None,
                "bbox": bbox_list,
                "current_tag": None,
            }
            pages_map[page_no].append(element)

        # ── IMAGE elements ─────────────────────────────────────────────────
        elif label == DocItemLabel.PICTURE:
            # Determine has_alt_text from struct tree; fall back to False.
            img_idx = _img_counters.get(page_no, 0)
            page_figures = _alt_map.get(page_no, [])
            has_alt_text: bool = page_figures[img_idx] if img_idx < len(page_figures) else False
            _img_counters[page_no] = img_idx + 1

            element = {
                "id": el_id,
                "type": "image",
                "docling_label": label.value,
                "bbox": bbox_list,
                "has_alt_text": has_alt_text,
            }
            pages_map[page_no].append(element)

        # ── TABLE elements ─────────────────────────────────────────────────
        elif label == DocItemLabel.TABLE:
            rows: int | None = None
            cols: int | None = None

            # docling TableItem exposes .data with num_rows / num_cols
            table_data = getattr(item, "data", None)
            if table_data is not None:
                rows = getattr(table_data, "num_rows", None)
                cols = getattr(table_data, "num_cols", None)

            # Extract cell content from the docling TableData grid
            cells: list[list[str]] | None = None
            if table_data is not None:
                raw_grid = getattr(table_data, "grid", None)
                if raw_grid:
                    try:
                        cells = []
                        for row in raw_grid:
                            cells.append([
                                cell.text if hasattr(cell, "text") else ""
                                for cell in row
                            ])
                    except Exception:
                        cells = None

            element = {
                "id": el_id,
                "type": "table",
                "docling_label": label.value,
                "rows": rows,
                "cols": cols,
                "has_header_row": "unknown",  # Claude infers this in the analysis layer
                "table_index": _table_counter,
                "cells": cells,
            }
            _table_counter += 1
            pages_map[page_no].append(element)

        # All other label types are skipped for now.

    # ── assemble pages list ────────────────────────────────────────────────
    pages = [
        {"page_number": pn, "elements": pages_map[pn]}
        for pn in sorted(pages_map.keys())
    ]

    return {
        "file_type": "pdf",
        "has_tag_tree": is_tagged_pdf(str(pdf_path)),
        "metadata": metadata,
        "pages": pages,
    }


def extract_docx(file_path: str | Path) -> dict:
    """
    Extract structured content from a Word .docx file.

    Returns the same JSON-serialisable dict schema as extract():
      {
        "file_type": "docx",
        "metadata": {"title": ..., "language": ..., "page_count": 1},
        "pages": [{"page_number": 1, "elements": [...]}]
      }
    """
    from docx import Document as DocxDocument

    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    doc = DocxDocument(str(file_path))

    # ── metadata ──────────────────────────────────────────────────────────
    props = doc.core_properties
    title: str | None = props.title or None
    language: str | None = props.language or None

    metadata = {
        "title": title,
        "language": language,
        "page_count": 1,  # docx has no reliable page count without rendering
    }

    _STYLE_MAP = {
        "Heading 1": "title",
        "Heading 2": "section_header",
        "Heading 3": "section_header",
        "Heading 4": "section_header",
        "List Paragraph": "list_item",
        "Caption": "caption",
    }

    elements: list[dict] = []
    counter = 1

    # ── paragraphs ────────────────────────────────────────────────────────
    for para_idx, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            continue

        style_name = para.style.name if para.style else "Normal"
        docling_label = _STYLE_MAP.get(style_name, "text")

        font_bold: bool | None = None
        if para.runs:
            font_bold = para.runs[0].bold

        # Check for inline images in this paragraph
        has_image = False
        for run in para.runs:
            xml = run._r.xml
            if "<a:blip" in xml or "drawing" in xml or "pic:" in xml:
                has_image = True
                break

        if has_image:
            # Detect alt text on drawing element
            has_alt = False
            try:
                from lxml import etree  # noqa: F401
                _WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                for run in para.runs:
                    for drawing in run._r.findall(f"{{{_WP_NS}}}docPr"):
                        descr = drawing.get("descr", "")
                        if descr and descr.strip():
                            has_alt = True
            except Exception:
                pass

            elements.append({
                "id": f"el_{counter:03d}",
                "type": "image",
                "docling_label": "picture",
                "text": None,
                "page": 1,
                "bbox": None,
                "current_tag": None,
                "has_alt_text": has_alt,
                "paragraph_index": para_idx,
            })
            counter += 1
        else:
            elements.append({
                "id": f"el_{counter:03d}",
                "type": "text",
                "docling_label": docling_label,
                "text": para.text,
                "page": 1,
                "bbox": None,
                "current_tag": style_name,
                "has_alt_text": None,
                "font_size": None,
                "font_bold": font_bold,
                "paragraph_index": para_idx,
            })
            counter += 1

    # ── tables ────────────────────────────────────────────────────────────
    for table_idx, table in enumerate(doc.tables):
        has_header = "unknown"
        if table.rows:
            first_cell_style = ""
            try:
                first_cell_style = table.rows[0].cells[0].paragraphs[0].style.name
            except (IndexError, AttributeError):
                pass
            if first_cell_style in ("Table Header", "Heading 1", "Heading 2"):
                has_header = "true"

        try:
            cells: list[list[str]] | None = [
                [cell.text for cell in row.cells]
                for row in table.rows
            ]
        except Exception:
            cells = None

        elements.append({
            "id": f"el_{counter:03d}",
            "type": "table",
            "docling_label": "table",
            "text": None,
            "page": 1,
            "bbox": None,
            "current_tag": None,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "has_header_row": has_header,
            "table_index": table_idx,
            "cells": cells,
        })
        counter += 1

    return {
        "file_type": "docx",
        "metadata": metadata,
        "pages": [{"page_number": 1, "elements": elements}],
    }


def is_tagged_pdf(pdf_path: str) -> bool:
    """
    Return True if the PDF has an accessibility tag tree (/StructTreeRoot).

    Parameters
    ----------
    pdf_path:
        Path to the PDF file.

    Returns
    -------
    bool — True if tagged, False if untagged or on any error.
    """
    try:
        import pikepdf
        with pikepdf.open(str(pdf_path)) as pdf:
            struct_tree = pdf.Root.get("/StructTreeRoot")
            return struct_tree is not None
    except Exception:
        return False


def render_element_thumbnail(
    pdf_path: str,
    page_number: int,
    bbox: list,
    max_px: int = 200,
) -> bytes:
    """
    Render a cropped region of a PDF page as PNG bytes.

    Parameters
    ----------
    pdf_path:
        Path to the PDF file.
    page_number:
        1-indexed page number.
    bbox:
        Bounding box as [x0, y0, x1, y1] in PDF point coordinates.
    max_px:
        Maximum size in pixels on the longest side before resizing.
        Display size is controlled separately by the caller.

    Returns
    -------
    bytes — PNG-encoded image data.
    """
    import io

    import fitz
    from PIL import Image as PILImage

    doc = fitz.open(pdf_path)
    page = doc[page_number - 1]  # fitz is 0-indexed
    page_height = page.rect.height
    x0, y0, x1, y1 = _docling_bbox_to_pymupdf(bbox, page_height)
    clip = fitz.Rect(x0, y0, x1, y1)
    pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), clip=clip)  # 2× render scale
    png_bytes = pix.tobytes("png")
    doc.close()

    # Resize to max_px * 2 on longest side using Pillow
    img = PILImage.open(io.BytesIO(png_bytes))
    img.thumbnail((max_px * 2, max_px * 2))
    out = io.BytesIO()
    img.save(out, format="PNG")
    return out.getvalue()


def extract_to_json(pdf_path: str | Path, output_path: str | Path | None = None) -> str:
    """
    Extract a PDF and return the result as a JSON string.

    Optionally writes the JSON to *output_path* as well.
    """
    result = extract(pdf_path)
    json_str = json.dumps(result, indent=2, default=str)
    if output_path is not None:
        Path(output_path).write_text(json_str, encoding="utf-8")
    return json_str
